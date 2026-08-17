import json
import subprocess
import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SCRIPT = ROOT / "scripts" / "progressive_papercheck.py"


def make_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("PaperCheck mode smoke", level=1)
    doc.add_paragraph("Transformer encoders are used in retrieval tasks [1].")
    doc.add_paragraph("This citation has no reference entry [2].")
    doc.add_heading("参考文献", level=1)
    doc.add_paragraph("[1] Vaswani A, Shazeer N, Parmar N, et al. Attention is all you need[C]. NeurIPS, 2017.")
    doc.add_paragraph("[3] Devlin J, Chang M W, Lee K, Toutanova K. BERT[C]. NAACL, 2019.")
    doc.save(path)


def test_progressive_papercheck_modes_and_report(tmp_path):
    paper = tmp_path / "paper.docx"
    evidence = tmp_path / "evidence.json"
    report = tmp_path / "report.md"
    make_docx(paper)

    completed = subprocess.run(
        [
            sys.executable,
            str(SCRIPT),
            str(paper),
            "--mode",
            "full",
            "--out",
            str(evidence),
            "--report",
            str(report),
        ],
        check=False,
        capture_output=True,
        text=True,
        encoding="utf-8",
    )

    assert completed.returncode == 0, completed.stderr
    events = [json.loads(line) for line in completed.stdout.splitlines() if line.strip()]
    event_names = [event["event"] for event in events]

    assert "mode_selected" in event_names
    assert "rules_check_complete" in event_names
    assert "model_review_ready" in event_names
    assert "source_verification_ready" in event_names
    assert "report_ready" in event_names
    assert events[-1]["event"] == "papercheck_complete"
    assert events[-1]["report"] == str(report)
    assert all(event.get("message") for event in events)
    assert any(event.get("message") == "报告已生成" for event in events)

    evidence_payload = json.loads(evidence.read_text(encoding="utf-8"))
    assert evidence_payload["citation_count"] == 2
    assert evidence_payload["missing_citations"] == ["[2]"]
    assert evidence_payload["unused_references"] == ["[3]"]

    markdown = report.read_text(encoding="utf-8")
    assert "Mode: `full`" in markdown
    assert "## Rules Summary" in markdown
    assert "## Current-Model Review Queue" in markdown
    assert "## Full-Mode Limits" in markdown
